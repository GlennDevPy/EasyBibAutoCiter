import requests
import json
import re
import os
from docx import Document
from docx.shared import Inches
from time import localtime, strftime
import codecs
import calendar





s = requests.Session()
os.system('cls')

bfind = s.get("https://www.easybib.com/api/auth/token?client=cfe", headers={
    "accept": "application/json",
    "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.146 Safari/537.36"
})
bearer = re.search('jwt\":\"(.*?)\"', bfind.text)[1]

cP = "https://bff.writing.chegg.com/graphql"
headersCP = {
    "accept": "*/*",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.146 Safari/537.36",
    "content-type": "application/json",
    "Referer": "https://www.easybib.com/",
    "Authorization": f"Bearer {bearer}",
    "Sec-GPC": "1",
    "Origin": "https://www.easybib.com"
}
headersMain = {
    "X-AlCell": "change-refresh-holdout-70-test",
    "Authorization": "Bearer anuZaem3ChaichuGhoh8gaiNeequee",
    "Origin": "https://www.easybib.com",
    "Referer": "https://www.easybib.com",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.101 Safari/537.36",
    "Accept": "application/vnd.com.easybib.data+json",
    "Sec-GPC": "1"
}


def createProj():
    data = {"operationName":"FoldersCreateProject","variables":{"folderId":None,"name":"EZBibCiter","defaultStyle":"mla8","public":True},"query":"mutation FoldersCreateProject($name: String!, $public: Boolean!, $defaultStyle: String!, $folderId: Int = null) {\n  createProject(name: $name, public: $public, defaultStyle: $defaultStyle, folderId: $folderId) {\n    id\n    name\n    defaultStyle\n    date: updatedAt\n    public\n    __typename\n  }\n}\n"}
    cPR = s.post(cP, headers=headersCP, data=json.dumps(data))
    ped = json.loads(cPR.text)
    bruh = (ped['data']['createProject']['id'])
    return bruh

def singleL(l1):
    worddoc = Document()
    sections = worddoc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    projID = createProj()
    singlelink = f"https://autocite.citation-api.com/api/v3/query?url={l1}"
    singler = s.get(singlelink, headers=headersMain, timeout=10)
    try:
        score = re.search('score\":(.*?),', singler.text)[1]
    except:
        score = 'Missing'

    try:
        title2 = re.search('title\":\"(.*?)\"', singler.text)[1]
        title = codecs.decode(title2, 'unicode-escape')
        if title == '':
            title = "Missing"
    except:
        title = 'Missing'

    try:
        webtitle = re.search('_mediumTitle\":\"(.*?)\"', singler.text)[1]
        if webtitle == '':
            webtitle = 'Missing'
    except:
        webtitle = 'Missing'

    try:
        authorF = re.search('given\":\"(.*?)\"', singler.text)[1]
        if authorF == '':
            authorF = 'Missing'
    except:
        authorF = 'Missing'

    try:
        authorL = re.search('family\":\"(.*?)\"', singler.text)[1]
        if authorL == '':
            authorL = "Missing"
    except:
        authorL = 'Missing'


    dD = re.search('_publishedDay\":\"(.*?)\"', singler.text)[1]
    if dD == '':
        dD = "Missing"


    dM = re.search('_publishedMonth\":\"(.*?)\"', singler.text)[1]
    if dM == '':
        dM = "Missing"


    dY = re.search('_publishedYear\":\"(.*?)\"', singler.text)[1]
    if dY == '':
        dY = "Missing"

    checklist = {
        "title": title,
        "webtitle": webtitle,
        "authorF": authorF,
        "authorL": authorL,
        "dD": dD,
        "dM": dM,
        "dY": dY,
    }

    tocheck = ['mediumTitleFound', 'contentTitleFound', 'contributorsFound', 'publishedDayFound', 'publishedMonthFound', 'publishedYearFound']
    #medtitle is website title
    if score == "100":
        print(f"\nCurrent Citation Progress:\n{checklist}\n")
        print("Do you want to make any edits? (Y or N)")
        cedits = input()
        if cedits.upper() == "Y":
            while True:
                print(f"URL: {l1}")
                print("\nTo which section?\n[1] Title\n[2] Website Title\n[3] Author First\n[4] Author Last\n[5] Day Pub\n[6] Month Pub\n[7] Year Pub\n[8] Done")
                sectionedit = input()
                if sectionedit == "1":
                    print(f'Change title from [{title}] to what?')
                    title = input()
                if sectionedit == "2":
                    print(f'Change website title from [{webtitle}] to what?')
                    webtitle = input()
                if sectionedit == "3":
                    print(f'Change author first name from [{authorF}] to what?')
                    authorF = input()
                if sectionedit == "4":
                    print(f'Change author last name from [{authorL}] to what?')
                    authorL = input()
                if sectionedit == "5":
                    print(f'Change day published to what?')
                    dD = input()
                if sectionedit == "6":
                    print(f'Change month published to what?')
                    dM = input()
                if sectionedit == "7":
                    print(f'Change year published to what?')
                    dY = input()
                if sectionedit == "8":
                    break
        else:
            pass
    if "Missing" in checklist.values():
        check = []
        for plz in tocheck:
            xPlz = re.search(f'{plz}\":(.*?),\"', singler.text)[1]
            if xPlz == "0":
                check.append(plz)
            else:
                pass
        
        print(f"\nCurrent Citation Progress:\n{checklist}\n")
        if "mediumTitleFound" in check:
            print("Website Title?")
            webtitle = input()

        if "contentTitleFound" in check:
            print("Article Title?")
            title = input()

        if "contributorsFound" in check:
            print("Author First Name?")
            authorF = input()

            print("Authors Last Name?")
            authorL = input()

        if "publishedDayFound" in check:
            print("Published Day? (DD)")
            dD = input()

        if "publishedMonthFound" in check:
            print("Published Month? (MM)")
            dM = input()
            
        if "publishedYearFound" in check:
            print("Published Year? (YYYY)")
            dY = input()
        
        checklist = {
        "title": title,
        "webtitle": webtitle,
        "authorF": authorF,
        "authorL": authorL,
        "dD": dD,
        "dM": dM,
        "dY": dY,
        }

        print(f"\nUpdated Citations:\n{checklist}\n")
        print("Do you want to make any edits? (Y or N)")
        cedits = input()
        if cedits.upper() == "Y":
            while True:
                print(f"URL: {l1}")
                print("\nTo which section?\n[1] Title\n[2] Website Title\n[3] Author First\n[4] Author Last\n[5] Day Pub\n[6] Month Pub\n[7] Year Pub\n[8] Done")
                sectionedit = input()
                if sectionedit == "1":
                    print('Change title to what?')
                    title = input()
                if sectionedit == "2":
                    print('Change website title to what?')
                    webtitle = input()
                if sectionedit == "3":
                    print('Change author first name to what?')
                    authorF = input()
                if sectionedit == "4":
                    print('Change author last name to what?')
                    authorL = input()
                if sectionedit == "5":
                    print('Change day published to what?')
                    dD = input()
                if sectionedit == "6":
                    print('Change month published to what?')
                    dM = input()
                if sectionedit == "7":
                    print('Change year published to what?')
                    dY = input()
                if sectionedit == "8":
                    break
        else:
            pass
    #print(f"\nUpdated Citations:\n{checklist}\n")
    #data2 = {"operationName":"CreateCitation","variables":{"projectId":projID,"contributors":[{"function":"author","first":authorF,"middle":"","last":authorL,"data":{"suffix":""}}],"annotation":"","data":{"source":"website","pubtype":{"main":"pubonline","suffix":""},"website":{"title":title},"pubonline":{"title":webtitle,"inst":"","url":l1,"day":dD,"month":dM,"year":dY,"timestamp":""},"annotation":"","validatorStatus":"complete"}},"query":"mutation CreateCitation($projectId: String!, $pubType: String, $sourceType: String, $annotation: String, $contributors: [ContributorInput!], $data: JSON!) {\n  createCitation(projectId: $projectId, pubType: $pubType, sourceType: $sourceType, annotation: $annotation, contributors: $contributors, data: $data) {\n    citationId\n    id: citationId\n    data\n    annotation\n    __typename\n  }\n}\n"}
    #cSC = s.post(cP, headers=headersCP, data=json.dumps(data2))
    #print(cSC.text)
    dM = dM.replace("0", "")
    if dM == '':
        paragraph = worddoc.add_paragraph(f'{authorL}, {authorF}. "{title}." {webtitle}, {dD}. {dY}, {l1}.')
    else:    
        paragraph = worddoc.add_paragraph(f'{authorL}, {authorF}. "{title}." {webtitle}, {dD} {calendar.month_abbr[int(dM)]}. {dY}, {l1}.')
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    worddoc.save(f'{strftime("%Y-%m-%d %H.%M.%S", localtime())} Citations.docx')
    print(f"\nFinished Citations - File at {strftime('%Y-%m-%d %H.%M.%S', localtime())} Citations.docx")

def textfile(tpath):
    os.system('cls')
    print("EZBibCiter [beta] by GlennDev\n")
    allcites = []
    worddoc = Document()
    sections = worddoc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    projID = createProj()
    links = []
    with open(tpath, 'r') as filepath:
        for i in filepath:
            links.append(i.rstrip())
    for letsgo in links:
        multilink = f"https://autocite.citation-api.com/api/v3/query?url={letsgo}"
        singler = s.get(multilink, headers=headersMain, timeout=10)
        try:
            score = re.search('score\":(.*?),', singler.text)[1]
        except:
            score = 'Missing'

        try:
            title2 = re.search('title\":\"(.*?)\"', singler.text)[1]
            title = codecs.decode(title2, 'unicode-escape')
        except:
            title = 'Missing'

        try:
            webtitle = re.search('_mediumTitle\":\"(.*?)\"', singler.text)[1]
        except:
            webtitle = 'Missing'

        try:
            authorF = re.search('given\":\"(.*?)\"', singler.text)[1]
        except:
            authorF = 'Missing'

        try:
            authorL = re.search('family\":\"(.*?)\"', singler.text)[1]
        except:
            authorL = 'Missing'


        dD = re.search('_publishedDay\":\"(.*?)\"', singler.text)[1]
        if dD == '':
            dD = "Missing"


        dM = re.search('_publishedMonth\":\"(.*?)\"', singler.text)[1]
        if dM == '':
            dM = "Missing"


        dY = re.search('_publishedYear\":\"(.*?)\"', singler.text)[1]
        if dY == '':
            dY = "Missing"
        

        checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            "dM": dM,
            "dY": dY,
            'furl': letsgo
        }

        tocheck = ['mediumTitleFound', 'contentTitleFound', 'contributorsFound', 'publishedDayFound', 'publishedMonthFound', 'publishedYearFound']
        
        #medtitle is website title
        if score == "100":
            print(f"\nCurrent Citation Progress:\n{checklist}\n")
            print("Do you want to make any edits? (Y or N)")
            cedits = input()
            if cedits.upper() == "Y":
                while True:
                    print(f"URL: {letsgo}")
                    print("\nTo which section?\n[1] Title\n[2] Website Title\n[3] Author First\n[4] Author Last\n[5] Day Pub\n[6] Month Pub\n[7] Year Pub\n[8] Done")
                    sectionedit = input()
                    if sectionedit == "1":
                        print(f'Change title from [{title}] to what?')
                        title = input()
                    if sectionedit == "2":
                        print(f'Change website title from [{webtitle}] to what?')
                        webtitle = input()
                    if sectionedit == "3":
                        print(f'Change author first name from [{authorF}] to what?')
                        authorF = input()
                    if sectionedit == "4":
                        print(f'Change author last name from [{authorL}] to what?')
                        authorL = input()
                    if sectionedit == "5":
                        print(f'Change day published to what?')
                        dD = input()
                    if sectionedit == "6":
                        print(f'Change month published to what?')
                        dM = input()
                    if sectionedit == "7":
                        print(f'Change year published to what?')
                        dY = input()
                    if sectionedit == "8":
                        break
            else:
                pass
        else:
            check = []
            for plz in tocheck:
                xPlz = re.search(f'{plz}\":(.*?),\"', singler.text)[1]
                if xPlz == "0":
                    check.append(plz)
                else:
                    pass
            
            print(f"\nCurrent Citation Progress:\n{checklist}\n")
            if "mediumTitleFound" in check:
                print("Website Title?")
                webtitle = input()

            if "contentTitleFound" in check:
                print("Article Title?")
                title = input()

            if "contributorsFound" in check:
                print("Author First Name?")
                authorF = input()

                print("Authors Last Name?")
                authorL = input()

            if "publishedDayFound" in check:
                print("Published Day? (DD)")
                dD = input()

            if "publishedMonthFound" in check:
                print("Published Month? (MM)")
                dM = input()
                
            if "publishedYearFound" in check:
                print("Published Year? (YYYY)")
                dY = input()
            
            checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            "dM": dM,
            "dY": dY,
            'furl': letsgo
            }

            print(f"\nUpdated Citations:\n{checklist}\n")
            print("Do you want to make any edits? (Y or N)")
            cedits = input()
            if cedits.upper() == "Y":
                while True:
                    print(f"URL: {letsgo}")
                    print("\nTo which section?\n[1] Title\n[2] Website Title\n[3] Author First\n[4] Author Last\n[5] Day Pub\n[6] Month Pub\n[7] Year Pub\n[8] Done")
                    sectionedit = input()
                    if sectionedit == "1":
                        print('Change title to what?')
                        title = input()
                    if sectionedit == "2":
                        print('Change website title to what?')
                        webtitle = input()
                    if sectionedit == "3":
                        print('Change author first name to what?')
                        authorF = input()
                    if sectionedit == "4":
                        print('Change author last name to what?')
                        authorL = input()
                    if sectionedit == "5":
                        print('Change day published to what?')
                        dD = input()
                    if sectionedit == "6":
                        print('Change month published to what?')
                        dM = input()
                    if sectionedit == "7":
                        print('Change year published to what?')
                        dY = input()
                    if sectionedit == "8":
                        break
            else:
                pass
        checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            "dM": dM,
            "dY": dY,
            'furl': letsgo
            }
        allcites.append(checklist)

        #print(f"\nUpdated Citations:\n{checklist}\n")
        #data2 = {"operationName":"CreateCitation","variables":{"projectId":projID,"contributors":[{"function":"author","first":authorF,"middle":"","last":authorL,"data":{"suffix":""}}],"annotation":"","data":{"source":"website","pubtype":{"main":"pubonline","suffix":""},"website":{"title":title},"pubonline":{"title":webtitle,"inst":"","url":letsgo,"day":dD,"month":dM,"year":dY,"timestamp":""},"annotation":"","validatorStatus":"complete"}},"query":"mutation CreateCitation($projectId: String!, $pubType: String, $sourceType: String, $annotation: String, $contributors: [ContributorInput!], $data: JSON!) {\n  createCitation(projectId: $projectId, pubType: $pubType, sourceType: $sourceType, annotation: $annotation, contributors: $contributors, data: $data) {\n    citationId\n    id: citationId\n    data\n    annotation\n    __typename\n  }\n}\n"}
        #cSC = s.post(cP, headers=headersCP, data=json.dumps(data2))

    while True:
        print(f"\nCurrent Citation List:\n")
        for index, item in enumerate(allcites, start=1):
            print(f"{index} | {item}")
        print(f'\n[1] Remove\n[2] Done')
        selectxxx3 = input()
        if selectxxx3 == "1":
            print("Which Citation Index to Remove? (1, 2, 3, etc)")
            indexrem = input()
            indexrem = int(indexrem)-1
            allcites.pop(indexrem)
        if selectxxx3 == "2":
            break
    sort = []
    final = []
    for values in allcites:
        sort.append(values["authorL"])
    sort.sort()
    for values2 in sort:
        #values 2 = sorted authors
        for values3 in allcites:
            #values 3 = original authors
            if values2 in values3.values():
                final.append(values3)
    #todo if author blank etc
    for outputs in final:
        outputs['dM'] = outputs['dM'].replace("0", "")
        if outputs['dM'] == '':
            paragraph = worddoc.add_paragraph(f"{outputs['authorL']}, {outputs['authorF']}. \"{outputs['title']}.\" {outputs['webtitle']}, {outputs['dD']}. {outputs['dY']}, {outputs['furl']}.")
        else:
            paragraph = worddoc.add_paragraph(f"{outputs['authorL']}, {outputs['authorF']}. \"{outputs['title']}.\" {outputs['webtitle']}, {outputs['dD']} {calendar.month_abbr[int(outputs['dM'])]}. {outputs['dY']}, {outputs['furl']}.")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    worddoc.save(f'{strftime("%Y-%m-%d %H.%M.%S", localtime())} Citations.docx')
    print(f"\nFinished Citations - File at {strftime('%Y-%m-%d %H.%M.%S', localtime())} Citations.docx")  
    #data3 = {"operationName":"Citationcount","variables":{"projectId":projID},"query":"query Citationcount($projectId: String!) {\n  citationCount(projectId: $projectId) {\n    count\n    __typename\n  }\n}\n"}
    #cx = s.post(cP, headers=headersCP, data=json.dumps(data3))
    #print(cx.text)
    #get the citations
    #dump this into word doc

def manual():
    os.system('cls')
    print("EZBibCiter [beta] by GlennDev\n")
    allcites = []
    worddoc = Document()
    sections = worddoc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    urls = []
    while True:
        print(f"Current URL List:\n")
        for index, item in enumerate(urls, start=1):
            print(f"{index} | {item}")
        print(f'\n[1] Add\n[2] Remove\n[3] Done')
        selectxxx3 = input()
        if selectxxx3 == "1":
            print("Enter URL to Add:\n")
            urladd = input()
            urls.append(urladd.rstrip())
            os.system('cls')
        if selectxxx3 == "2":
            print("Which URL Index to Remove? (1, 2, 3, etc)")
            indexrem = input()
            indexrem = int(indexrem)-1
            urls.pop(indexrem)
        if selectxxx3 == "3":
            break
    for letsgo in urls:
        multilink = f"https://autocite.citation-api.com/api/v3/query?url={letsgo}"
        singler = s.get(multilink, headers=headersMain, timeout=10)
        try:
            score = re.search('score\":(.*?),', singler.text)[1]
        except:
            score = 'Missing'

        try:
            title2 = re.search('title\":\"(.*?)\"', singler.text)[1]
            title = codecs.decode(title2, 'unicode-escape')
        except:
            title = 'Missing'

        try:
            webtitle = re.search('_mediumTitle\":\"(.*?)\"', singler.text)[1]
        except:
            webtitle = 'Missing'

        try:
            authorF = re.search('given\":\"(.*?)\"', singler.text)[1]
        except:
            authorF = 'Missing'

        try:
            authorL = re.search('family\":\"(.*?)\"', singler.text)[1]
        except:
            authorL = 'Missing'


        dD = re.search('_publishedDay\":\"(.*?)\"', singler.text)[1]
        if dD == '':
            dD = "Missing"


        dM = re.search('_publishedMonth\":\"(.*?)\"', singler.text)[1]
        if dM == '':
            dM = "Missing"


        dY = re.search('_publishedYear\":\"(.*?)\"', singler.text)[1]
        if dY == '':
            dY = "Missing"
        

        checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            
            "dM": dM,
            "dY": dY,
            'furl': letsgo
        }

        tocheck = ['mediumTitleFound', 'contentTitleFound', 'contributorsFound', 'publishedDayFound', 'publishedMonthFound', 'publishedYearFound']
        
        #medtitle is website title
        if score == "100":
            print(f"\nCurrent Citation Progress:\n{checklist}\n")
            print("Do you want to make any edits? (Y or N)")
            cedits = input()
            if cedits.upper() == "Y":
                while True:
                    print(f"URL: {letsgo}")
                    print("\nTo which section?\n[1] Title\n[2] Website Title\n[3] Author First\n[4] Author Last\n[5] Day Pub\n[6] Month Pub\n[7] Year Pub\n[8] Done")
                    sectionedit = input()
                    if sectionedit == "1":
                        print(f'Change title from [{title}] to what?')
                        title = input()
                    if sectionedit == "2":
                        print(f'Change website title from [{webtitle}] to what?')
                        webtitle = input()
                    if sectionedit == "3":
                        print(f'Change author first name from [{authorF}] to what?')
                        authorF = input()
                    if sectionedit == "4":
                        print(f'Change author last name from [{authorL}] to what?')
                        authorL = input()
                    if sectionedit == "5":
                        print(f'Change day published to what?')
                        dD = input()
                    if sectionedit == "6":
                        print(f'Change month published to what?')
                        dM = input()
                    if sectionedit == "7":
                        print(f'Change year published to what?')
                        dY = input()
                    if sectionedit == "8":
                        break
            else:
                pass
        else:
            check = []
            for plz in tocheck:
                xPlz = re.search(f'{plz}\":(.*?),\"', singler.text)[1]
                if xPlz == "0":
                    check.append(plz)
                else:
                    pass
            
            print(f"\nCurrent Citation Progress:\n{checklist}\n")
            if "mediumTitleFound" in check:
                print("Website Title?")
                webtitle = input()

            if "contentTitleFound" in check:
                print("Article Title?")
                title = input()

            if "contributorsFound" in check:
                print("Author First Name?")
                authorF = input()

                print("Authors Last Name?")
                authorL = input()

            if "publishedDayFound" in check:
                print("Published Day? (YY)")
                dD = input()

            if "publishedMonthFound" in check:
                print("Published Month? (MM)")
                dM = input()
                
            if "publishedYearFound" in check:
                print("Published Year? (YYYY)")
                dY = input()
            
            checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            "dM": dM,
            "dY": dY,
            'furl': letsgo
            }

            print(f"\nUpdated Citations:\n{checklist}\n")
            print("Do you want to make any edits? (Y or N)")
            cedits = input()
            if cedits.upper() == "Y":
                while True:
                    print(f"URL: {letsgo}")
                    print("\nTo which section?\n[1] Title\n[2] Website Title\n[3] Author First\n[4] Author Last\n[5] Day Pub\n[6] Month Pub\n[7] Year Pub\n[8] Done")
                    sectionedit = input()
                    if sectionedit == "1":
                        print('Change title to what?')
                        title = input()
                    if sectionedit == "2":
                        print('Change website title to what?')
                        webtitle = input()
                    if sectionedit == "3":
                        print('Change author first name to what?')
                        authorF = input()
                    if sectionedit == "4":
                        print('Change author last name to what?')
                        authorL = input()
                    if sectionedit == "5":
                        print('Change day published to what?')
                        dD = input()
                    if sectionedit == "6":
                        print('Change month published to what?')
                        dM = input()
                    if sectionedit == "7":
                        print('Change year published to what?')
                        dY = input()
                    if sectionedit == "8":
                        break
            else:
                pass
        checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            "dM": dM,
            "dY": dY,
            'furl': letsgo
        }
        allcites.append(checklist)
        #print(f"\nUpdated Citations:\n{checklist}\n")
        #data2 = {"operationName":"CreateCitation","variables":{"projectId":projID,"contributors":[{"function":"author","first":authorF,"middle":"","last":authorL,"data":{"suffix":""}}],"annotation":"","data":{"source":"website","pubtype":{"main":"pubonline","suffix":""},"website":{"title":title},"pubonline":{"title":webtitle,"inst":"","url":letsgo,"day":dD,"month":dM,"year":dY,"timestamp":""},"annotation":"","validatorStatus":"complete"}},"query":"mutation CreateCitation($projectId: String!, $pubType: String, $sourceType: String, $annotation: String, $contributors: [ContributorInput!], $data: JSON!) {\n  createCitation(projectId: $projectId, pubType: $pubType, sourceType: $sourceType, annotation: $annotation, contributors: $contributors, data: $data) {\n    citationId\n    id: citationId\n    data\n    annotation\n    __typename\n  }\n}\n"}
        #cSC = s.post(cP, headers=headersCP, data=json.dumps(data2))
    sort = []
    final = []
    for values in allcites:
        sort.append(values["authorL"])
    sort.sort()
    for values2 in sort:
        #values 2 = sorted authors
        for values3 in allcites:
            #values 3 = original authors
            if values2 in values3.values():
                final.append(values3)
    for outputs in final:
        outputs['dM'] = outputs['dM'].replace("0", "")
        paragraph = worddoc.add_paragraph(f"{outputs['authorL']}, {outputs['authorF']}. \"{outputs['title']}.\" {outputs['webtitle']}, {outputs['dD']} {calendar.month_abbr[int(outputs['dM'])]}. {outputs['dY']}, {outputs['furl']}.")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    worddoc.save(f'{strftime("%Y-%m-%d %H.%M.%S", localtime())} Citations.docx')
    print(f"\nFinished Citations - File at {strftime('%Y-%m-%d %H.%M.%S', localtime())} Citations.docx")
    
    #data3 = {"operationName":"Citationcount","variables":{"projectId":projID},"query":"query Citationcount($projectId: String!) {\n  citationCount(projectId: $projectId) {\n    count\n    __typename\n  }\n}\n"}
    #cx = s.post(cP, headers=headersCP, data=json.dumps(data3))
    #print(cx.text)
    #get the citations
    #dump this into word doc

def noexppath():
    os.system('cls')
    print("EZBibCiter [beta] by GlennDev\n")
    allcites = []
    urls = []
    while True:
        print(f"Current URL List:\n")
        for index, item in enumerate(urls, start=1):
            print(f"{index} | {item}")
        print(f'\n[1] Add\n[2] Remove\n[3] Done')
        selectxxx3 = input()
        if selectxxx3 == "1":
            print("Enter URL to Add:\n")
            urladd = input()
            urls.append(urladd.rstrip())
            os.system('cls')
        if selectxxx3 == "2":
            print("Which URL Index to Remove? (1, 2, 3, etc)")
            indexrem = input()
            indexrem = int(indexrem)-1
            urls.pop(indexrem)
        if selectxxx3 == "3":
            break
    for letsgo in urls:
        multilink = f"https://autocite.citation-api.com/api/v3/query?url={letsgo}"
        singler = s.get(multilink, headers=headersMain, timeout=10)
        try:
            score = re.search('score\":(.*?),', singler.text)[1]
        except:
            score = 'Missing'

        try:
            title2 = re.search('title\":\"(.*?)\"', singler.text)[1]
            title = codecs.decode(title2, 'unicode-escape')
        except:
            title = 'Missing'

        try:
            webtitle = re.search('_mediumTitle\":\"(.*?)\"', singler.text)[1]
        except:
            webtitle = 'Missing'

        try:
            authorF = re.search('given\":\"(.*?)\"', singler.text)[1]
        except:
            authorF = 'Missing'

        try:
            authorL = re.search('family\":\"(.*?)\"', singler.text)[1]
        except:
            authorL = 'Missing'


        dD = re.search('_publishedDay\":\"(.*?)\"', singler.text)[1]
        if dD == '':
            dD = "Missing"


        dM = re.search('_publishedMonth\":\"(.*?)\"', singler.text)[1]
        if dM == '':
            dM = "Missing"


        dY = re.search('_publishedYear\":\"(.*?)\"', singler.text)[1]
        if dY == '':
            dY = "Missing"
        

        checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            
            "dM": dM,
            "dY": dY,
            'furl': letsgo
        }

        tocheck = ['mediumTitleFound', 'contentTitleFound', 'contributorsFound', 'publishedDayFound', 'publishedMonthFound', 'publishedYearFound']
        
        #medtitle is website title
        if score == "100":
            print(f"\nCurrent Citation Progress:\n{checklist}\n")
            print("Do you want to make any edits? (Y or N)")
            cedits = input()
            if cedits.upper() == "Y":
                while True:
                    print(f"URL: {letsgo}")
                    print("\nTo which section?\n[1] Title\n[2] Website Title\n[3] Author First\n[4] Author Last\n[5] Day Pub\n[6] Month Pub\n[7] Year Pub\n[8] Done")
                    sectionedit = input()
                    if sectionedit == "1":
                        print(f'Change title from [{title}] to what?')
                        title = input()
                    if sectionedit == "2":
                        print(f'Change website title from [{webtitle}] to what?')
                        webtitle = input()
                    if sectionedit == "3":
                        print(f'Change author first name from [{authorF}] to what?')
                        authorF = input()
                    if sectionedit == "4":
                        print(f'Change author last name from [{authorL}] to what?')
                        authorL = input()
                    if sectionedit == "5":
                        print(f'Change day published to what?')
                        dD = input()
                    if sectionedit == "6":
                        print(f'Change month published to what?')
                        dM = input()
                    if sectionedit == "7":
                        print(f'Change year published to what?')
                        dY = input()
                    if sectionedit == "8":
                        break
            else:
                pass
        else:
            check = []
            for plz in tocheck:
                xPlz = re.search(f'{plz}\":(.*?),\"', singler.text)[1]
                if xPlz == "0":
                    check.append(plz)
                else:
                    pass
            
            print(f"\nCurrent Citation Progress:\n{checklist}\n")
            if "mediumTitleFound" in check:
                print("Website Title?")
                webtitle = input()

            if "contentTitleFound" in check:
                print("Article Title?")
                title = input()

            if "contributorsFound" in check:
                print("Author First Name?")
                authorF = input()

                print("Authors Last Name?")
                authorL = input()

            if "publishedDayFound" in check:
                print("Published Day? (YY)")
                dD = input()

            if "publishedMonthFound" in check:
                print("Published Month? (MM)")
                dM = input()
                
            if "publishedYearFound" in check:
                print("Published Year? (YYYY)")
                dY = input()
            
            checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            "dM": dM,
            "dY": dY,
            'furl': letsgo
            }

            print(f"\nUpdated Citations:\n{checklist}\n")
            print("Do you want to make any edits? (Y or N)")
            cedits = input()
            if cedits.upper() == "Y":
                while True:
                    print(f"URL: {letsgo}")
                    print("\nTo which section?\n[1] Title\n[2] Website Title\n[3] Author First\n[4] Author Last\n[5] Day Pub\n[6] Month Pub\n[7] Year Pub\n[8] Done")
                    sectionedit = input()
                    if sectionedit == "1":
                        print('Change title to what?')
                        title = input()
                    if sectionedit == "2":
                        print('Change website title to what?')
                        webtitle = input()
                    if sectionedit == "3":
                        print('Change author first name to what?')
                        authorF = input()
                    if sectionedit == "4":
                        print('Change author last name to what?')
                        authorL = input()
                    if sectionedit == "5":
                        print('Change day published to what?')
                        dD = input()
                    if sectionedit == "6":
                        print('Change month published to what?')
                        dM = input()
                    if sectionedit == "7":
                        print('Change year published to what?')
                        dY = input()
                    if sectionedit == "8":
                        break
            else:
                pass
        checklist = {
            "title": title,
            "webtitle": webtitle,
            "authorF": authorF,
            "authorL": authorL,
            "dD": dD,
            "dM": dM,
            "dY": dY,
            'furl': letsgo
        }
        allcites.append(checklist)
        #print(f"\nUpdated Citations:\n{checklist}\n")
        #data2 = {"operationName":"CreateCitation","variables":{"projectId":projID,"contributors":[{"function":"author","first":authorF,"middle":"","last":authorL,"data":{"suffix":""}}],"annotation":"","data":{"source":"website","pubtype":{"main":"pubonline","suffix":""},"website":{"title":title},"pubonline":{"title":webtitle,"inst":"","url":letsgo,"day":dD,"month":dM,"year":dY,"timestamp":""},"annotation":"","validatorStatus":"complete"}},"query":"mutation CreateCitation($projectId: String!, $pubType: String, $sourceType: String, $annotation: String, $contributors: [ContributorInput!], $data: JSON!) {\n  createCitation(projectId: $projectId, pubType: $pubType, sourceType: $sourceType, annotation: $annotation, contributors: $contributors, data: $data) {\n    citationId\n    id: citationId\n    data\n    annotation\n    __typename\n  }\n}\n"}
        #cSC = s.post(cP, headers=headersCP, data=json.dumps(data2))
    sort = []
    final = []
    for values in allcites:
        sort.append(values["authorL"])
    sort.sort()
    for values2 in sort:
        #values 2 = sorted authors
        for values3 in allcites:
            #values 3 = original authors
            if values2 in values3.values():
                final.append(values3)
    for outputs in final:
        outputs['dM'] = outputs['dM'].replace("0", "")
        paragraph = worddoc.add_paragraph(f"{outputs['authorL']}, {outputs['authorF']}. \"{outputs['title']}.\" {outputs['webtitle']}, {outputs['dD']} {calendar.month_abbr[int(outputs['dM'])]}. {outputs['dY']}, {outputs['furl']}.")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    worddoc.save(f'{strftime("%Y-%m-%d %H.%M.%S", localtime())} Citations.docx')
    print(f"\nFinished Citations - File at {strftime('%Y-%m-%d %H.%M.%S', localtime())} Citations.docx")
    
    #data3 = {"operationName":"Citationcount","variables":{"projectId":projID},"query":"query Citationcount($projectId: String!) {\n  citationCount(projectId: $projectId) {\n    count\n    __typename\n  }\n}\n"}
    #cx = s.post(cP, headers=headersCP, data=json.dumps(data3))
    #print(cx.text)
    #get the citations
    #dump this into word doc  


print("EZBibCiter [beta] by GlennDev\n")
print("[1] Textfile\n[2] Single Link\n[3] Enter Links Manually")
choice = input()

if choice == "1":
    print("\nDrag in Text File/Give Path to File")
    txtfile = input()
    textfile(txtfile)

if choice == "2":
    print("\nEnter in the link to cite:")
    xxbru = input()
    singleL(xxbru)


if choice == "3":
    manual()
    
if choice == "4":
    noexppath()